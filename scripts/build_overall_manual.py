from pathlib import Path
import re

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
MANUAL_MD = ROOT / "CODE_RECALL_FULL_SYSTEM_MANUAL_DRAFT.md"
OUT = ROOT / "manuals" / "Overall_Manual.docx"
COVER = ROOT / "manuals" / "overall-manual-cover.png"


OPENING_STATEMENT = [
    "The researchers present this manual as a complete guide for the use, management, and future development of the Code Recall system. This document was prepared to help users, administrators, super administrators, and developers understand the purpose, functions, structure, and proper operation of the system.",
    "Code Recall was developed as a gamified learning platform designed to support learners in studying Computer System Servicing-related topics through interactive lessons, assessments, progress tracking, achievements, and review tools. The system aims to provide a more engaging and organized learning experience by combining educational content with game-based elements such as XP, badges, certificates, and leaderboards.",
    "This manual serves as both an operational guide and a technical reference. It explains how learners can access the system, use the Play as Guest feature, complete subject activities, answer assessments, review progress, and unlock certificates. It also explains how administrators and super administrators can monitor learner activity, manage system-related responsibilities, and support the proper use of the platform.",
    "The researchers prepared this manual to ensure that the system can be used properly, maintained effectively, and improved continuously. It is intended to support system implementation, user training, administrator orientation, developer handover, and future enhancement of the Code Recall platform.",
    "Through this manual, the researchers hope that users will be guided clearly in using the system, administrators will be assisted in managing learner progress, and future developers will be able to understand and continue improving the project.",
]

RESEARCHER_MESSAGE = [
    "The researchers would like to present this manual as a guide for understanding and using the Code Recall system. This manual was created to support learners, administrators, super administrators, and future developers who will use, manage, maintain, or improve the system.",
    "Code Recall was developed with the goal of making learning more engaging, organized, and interactive. The system focuses on Computer System Servicing-related topics and uses gamified features such as XP, achievements, certificates, progress tracking, and leaderboards to encourage learners to participate actively in the learning process.",
    "The researchers believe that a system is more useful when its users understand how to operate it properly and when future maintainers understand how it was built. For this reason, this manual was prepared not only as a user guide, but also as a reference for system management and future development.",
    "It is the hope of the researchers that this manual will help users navigate the Code Recall system with confidence, assist administrators in monitoring learner progress effectively, and guide future developers in maintaining and enhancing the platform for continued use.",
]


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_borders(cell, color="CBD5E1", size="8"):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for side in ("top", "left", "bottom", "right"):
        el = borders.find(qn(f"w:{side}"))
        if el is None:
            el = OxmlElement(f"w:{side}")
            borders.append(el)
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), size)
        el.set(qn("w:color"), color)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Page ")
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.append(begin)
    run._r.append(instr)
    run._r.append(end)


def configure_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.08

    style_specs = {
        "Title": (22, RGBColor(30, 64, 175)),
        "Subtitle": (12, RGBColor(71, 85, 105)),
        "Heading 1": (16, RGBColor(30, 64, 175)),
        "Heading 2": (14, RGBColor(88, 28, 135)),
        "Heading 3": (12, RGBColor(15, 23, 42)),
    }
    for style_name, (size, color) in style_specs.items():
        style = styles[style_name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
        style.font.size = Pt(size)
        style.font.color.rgb = color
        if "Heading" in style_name or style_name == "Title":
            style.font.bold = True
        if style_name == "Heading 1":
            style.paragraph_format.page_break_before = False

    for list_style in ("List Bullet", "List Number"):
        style = styles[list_style]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
        style.font.size = Pt(11)
        style.paragraph_format.space_after = Pt(4)


def add_paragraph(doc, text="", style=None):
    paragraph = doc.add_paragraph(style=style)
    paragraph.add_run(text)
    return paragraph


def add_caption(doc, text):
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(text)
    run.italic = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(80, 80, 80)
    return paragraph


def add_screenshot_placeholder(doc, caption, figure_number):
    table = doc.add_table(rows=5, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for row_index, row in enumerate(table.rows):
        row.height = Inches(0.34)
        cell = row.cells[0]
        cell.width = Inches(6.2)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        shade_cell(cell, "F1F5F9")
        set_cell_borders(cell)
        if row_index == 2:
            cell.text = ""
            paragraph = cell.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = paragraph.add_run("INSERT SCREENSHOT HERE")
            run.bold = True
            run.font.size = Pt(12)
            run.font.color.rgb = RGBColor(71, 85, 105)
    add_caption(doc, f"Figure {figure_number}. {caption}")


def add_markdown_table(doc, table_lines):
    rows = []
    for line in table_lines:
        stripped = line.strip()
        if not stripped.startswith("|"):
            continue
        cells = [cell.strip() for cell in stripped.strip("|").split("|")]
        if all(re.fullmatch(r":?-{3,}:?", cell or "") for cell in cells):
            continue
        rows.append(cells)
    if not rows:
        return
    table = doc.add_table(rows=1, cols=len(rows[0]))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for col_index, value in enumerate(rows[0]):
        cell = table.rows[0].cells[col_index]
        shade_cell(cell, "E2E8F0")
        set_cell_borders(cell)
        cell.text = ""
        run = cell.paragraphs[0].add_run(value)
        run.bold = True
        run.font.size = Pt(9)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    for row in rows[1:]:
        cells = table.add_row().cells
        for col_index, value in enumerate(row):
            cell = cells[col_index]
            set_cell_borders(cell)
            cell.text = value
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    doc.add_paragraph()


def add_title_page(doc):
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run("CODE RECALL SYSTEM MANUAL")
    run.bold = True
    run.font.size = Pt(24)
    run.font.color.rgb = RGBColor(30, 64, 175)

    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run("User, Admin, Super Admin, and Developer Guide")
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(88, 28, 135)

    doc.add_paragraph()
    add_paragraph(doc, "Prepared for the Code Recall Gamified Learning System")
    add_paragraph(
        doc,
        "This manual documents the learner, guest, administrator, super administrator, and developer workflows of the Code Recall system. It explains how users operate the system, how administrators monitor learner activity, how super administrators manage privileged access, and how future developers may maintain and extend the project.",
    )

    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.rows[0].cells[0].text = "Document Field"
    table.rows[0].cells[1].text = "Description"
    for cell in table.rows[0].cells:
        shade_cell(cell, "E2E8F0")
        set_cell_borders(cell)
        for run in cell.paragraphs[0].runs:
            run.bold = True
    rows = [
        ("System Name", "Code Recall"),
        ("Document Type", "Comprehensive System Manual"),
        ("Primary Users", "Learners, Guests, Admins, Super Admins, Developers"),
        ("Prepared For", "Implementation, thesis documentation, user training, and developer handover"),
        ("Paper Size", "Short bond paper / US Letter, 8.5 x 11 inches"),
    ]
    for left, right in rows:
        cells = table.add_row().cells
        cells[0].text = left
        cells[1].text = right
        for cell in cells:
            set_cell_borders(cell)


def add_opening_sections(doc):
    doc.add_page_break()
    doc.add_heading("Opening Statement", level=1)
    for text in OPENING_STATEMENT:
        add_paragraph(doc, text)
    doc.add_heading("Message from the Researchers", level=1)
    for text in RESEARCHER_MESSAGE:
        add_paragraph(doc, text)


def prepare_markdown_lines():
    lines = MANUAL_MD.read_text(encoding="utf-8").splitlines()
    cleaned = []
    skip_until_first_section = False
    for line in lines:
        if line.strip() == "# Code Recall Full System Manual Draft":
            skip_until_first_section = True
            continue
        if skip_until_first_section:
            if line.strip() == "# Table of Contents":
                skip_until_first_section = False
                cleaned.append(line)
            continue
        if line.strip() in {"# Title Page", "## Code Recall System Manual"}:
            continue
        cleaned.append(line)
    return cleaned


def render_markdown(doc):
    figure_number = 1
    lines = prepare_markdown_lines()
    index = 0
    while index < len(lines):
        line = lines[index].rstrip()
        stripped = line.strip()
        if not stripped:
            index += 1
            continue
        if stripped == "---":
            doc.add_paragraph()
            index += 1
            continue
        if stripped.startswith("|"):
            table_lines = []
            while index < len(lines) and lines[index].strip().startswith("|"):
                table_lines.append(lines[index])
                index += 1
            add_markdown_table(doc, table_lines)
            continue
        if stripped.lower() == "suggested screenshot:":
            caption = "Suggested screenshot placeholder."
            if index + 2 < len(lines) and lines[index + 2].strip().startswith("- "):
                caption = lines[index + 2].strip()[2:].strip().rstrip(".") + "."
                index += 3
            add_screenshot_placeholder(doc, caption, figure_number)
            figure_number += 1
            continue
        if stripped.startswith("# "):
            heading = stripped[2:].strip()
            if heading in {"Table of Contents", "1. Introduction"}:
                doc.add_page_break()
            doc.add_heading(heading, level=1)
            index += 1
            continue
        if stripped.startswith("## "):
            doc.add_heading(stripped[3:].strip(), level=2)
            index += 1
            continue
        if stripped.startswith("### "):
            doc.add_heading(stripped[4:].strip(), level=3)
            index += 1
            continue
        if stripped.startswith("- "):
            add_paragraph(doc, stripped[2:].strip(), style="List Bullet")
            index += 1
            continue
        numbered = re.match(r"^\\d+\\.\\s+(.*)", stripped)
        if numbered:
            add_paragraph(doc, numbered.group(1), style="List Number")
            index += 1
            continue
        add_paragraph(doc, stripped.replace("`", ""))
        index += 1


def main():
    doc = Document()
    cover_section = doc.sections[0]
    cover_section.page_width = Inches(8.5)
    cover_section.page_height = Inches(11)
    cover_section.top_margin = Inches(0.35)
    cover_section.bottom_margin = Inches(0.35)
    cover_section.left_margin = Inches(0.4)
    cover_section.right_margin = Inches(0.4)
    configure_styles(doc)

    if COVER.exists():
        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.add_run().add_picture(str(COVER), width=Inches(7.55))
    else:
        doc.add_heading("CODE RECALL SYSTEM MANUAL", 0)

    doc.add_section(WD_SECTION.NEW_PAGE)
    body_section = doc.sections[-1]
    body_section.header.is_linked_to_previous = False
    body_section.footer.is_linked_to_previous = False
    body_section.page_width = Inches(8.5)
    body_section.page_height = Inches(11)
    body_section.top_margin = Inches(1)
    body_section.bottom_margin = Inches(1)
    body_section.left_margin = Inches(1)
    body_section.right_margin = Inches(1)
    header = body_section.header.paragraphs[0]
    header.text = "Code Recall System Manual"
    header.runs[0].font.size = Pt(9)
    header.runs[0].font.color.rgb = RGBColor(100, 116, 139)
    add_page_number(body_section.footer.paragraphs[0])

    add_title_page(doc)
    add_opening_sections(doc)
    render_markdown(doc)
    OUT.parent.mkdir(exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()

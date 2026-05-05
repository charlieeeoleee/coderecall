from __future__ import annotations

from dataclasses import dataclass
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent.parent
OUTPUT_DIR = ROOT / "manuals"
OUTPUT_DIR.mkdir(exist_ok=True)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color: str = "CFCBFF") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "8")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_table_cell_margins(cell, top=100, start=140, bottom=100, end=140) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def add_header(section, label: str) -> None:
    paragraph = section.header.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.space_after = Pt(6)
    run = paragraph.add_run(label)
    run.font.name = "Arial"
    run.font.size = Pt(9)
    run.font.bold = True
    run.font.color.rgb = RGBColor(88, 87, 136)

    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), "D9D7F5")
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def configure_document(document: Document, label: str) -> None:
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    add_header(section, label)

    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.08

    styles["Title"].font.name = "Arial"
    styles["Title"].font.size = Pt(24)
    styles["Title"].font.bold = True
    styles["Title"].font.color.rgb = RGBColor(31, 26, 77)

    styles["Subtitle"].font.name = "Arial"
    styles["Subtitle"].font.size = Pt(12)
    styles["Subtitle"].font.color.rgb = RGBColor(88, 87, 136)

    for name, size in (("Heading 1", 16), ("Heading 2", 13), ("Heading 3", 11)):
        style = styles[name]
        style.font.name = "Arial"
        style.font.bold = True
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor(31, 26, 77)


def add_title_page(document: Document, title: str, subtitle: str, audience: str) -> None:
    p = document.add_paragraph(style="Title")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(title)

    sub = document.add_paragraph(style="Subtitle")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.add_run(subtitle)

    meta = document.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.paragraph_format.space_before = Pt(8)
    meta.paragraph_format.space_after = Pt(18)
    meta.add_run(f"Audience: {audience}\n").bold = True
    meta.add_run(f"Prepared: {datetime.now().strftime('%B %d, %Y')}")

    callout = document.add_table(rows=1, cols=1)
    callout.autofit = False
    cell = callout.rows[0].cells[0]
    cell.width = Inches(6.3)
    set_cell_shading(cell, "F3F1FF")
    set_cell_border(cell, "D8D2FF")
    set_table_cell_margins(cell, 120, 180, 120, 180)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Document Note\n")
    r.bold = True
    r.font.size = Pt(12)
    r.font.color.rgb = RGBColor(31, 26, 77)
    body = p.add_run("This manual is designed to be screenshot-ready. Add the listed images after capturing the final live system screens.")
    body.font.size = Pt(10.5)
    body.font.color.rgb = RGBColor(88, 87, 136)

    document.add_paragraph()
    document.add_page_break()


def add_paragraph(document: Document, text: str) -> None:
    document.add_paragraph(text)


def add_bullets(document: Document, items: list[str]) -> None:
    for item in items:
        document.add_paragraph(item, style="List Bullet")


def add_numbered(document: Document, items: list[str]) -> None:
    for item in items:
        document.add_paragraph(item, style="List Number")


def add_screenshot_placeholder(document: Document, filename: str, description: str) -> None:
    document.add_paragraph("Suggested screenshot", style="Heading 3")
    table = document.add_table(rows=1, cols=1)
    table.autofit = False
    cell = table.rows[0].cells[0]
    cell.width = Inches(6.3)
    set_cell_shading(cell, "F7F6FF")
    set_cell_border(cell, "CEC8FF")
    set_table_cell_margins(cell, 180, 200, 180, 200)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = p.add_run("IMAGE PLACEHOLDER\n")
    title_run.bold = True
    title_run.font.size = Pt(11)
    title_run.font.color.rgb = RGBColor(31, 26, 77)
    file_run = p.add_run(f"{filename}\n")
    file_run.bold = True
    file_run.font.size = Pt(10)
    file_run.font.color.rgb = RGBColor(54, 112, 199)
    desc_run = p.add_run(description)
    desc_run.font.size = Pt(10)
    desc_run.font.color.rgb = RGBColor(88, 87, 136)
    document.add_paragraph()


@dataclass
class ManualSection:
    title: str
    paragraphs: list[str]
    bullets: list[str] | None = None
    numbered: list[str] | None = None
    screenshot_name: str | None = None
    screenshot_description: str | None = None


def build_manual(filename: str, title: str, subtitle: str, audience: str, header_label: str, sections: list[ManualSection]) -> Path:
    document = Document()
    configure_document(document, header_label)
    add_title_page(document, title, subtitle, audience)

    for section in sections:
        document.add_paragraph(section.title, style="Heading 1")
        for paragraph in section.paragraphs:
            add_paragraph(document, paragraph)
        if section.bullets:
            add_bullets(document, section.bullets)
        if section.numbered:
            add_numbered(document, section.numbered)
        if section.screenshot_name and section.screenshot_description:
            add_screenshot_placeholder(document, section.screenshot_name, section.screenshot_description)

    path = OUTPUT_DIR / filename
    document.save(path)
    return path


def build_checklist_doc(filename: str, title: str, rows: list[tuple[str, str, tuple[str, str]]]) -> Path:
    document = Document()
    configure_document(document, "Code Recall | Screenshot Checklist")
    add_title_page(document, title, "Image capture guide for all role-based manuals", "Documentation team")

    document.add_paragraph("Capture Instructions", style="Heading 1")
    document.add_paragraph("Use this checklist while taking screenshots for the final manuals. Keep all captures clear, complete, and consistent in theme.")
    add_bullets(document, [
        "Use the same theme whenever possible.",
        "Prefer full desktop captures instead of tight crops unless the section is intentionally focused.",
        "Make sure sample data is readable and acceptable for documentation.",
        "Capture final live screens only after the latest code is deployed or refreshed.",
    ])

    document.add_paragraph("Combined Screenshot Checklist", style="Heading 1")
    table = document.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    table.autofit = False
    widths = [Inches(0.7), Inches(1.55), Inches(1.75), Inches(2.4)]
    headers = ["No.", "Manual", "Filename", "What to capture"]
    for idx, cell in enumerate(table.rows[0].cells):
        cell.width = widths[idx]
        cell.text = headers[idx]
        set_cell_shading(cell, "ECE9FF")
        set_cell_border(cell, "CFC8FF")
        set_table_cell_margins(cell, 90, 120, 90, 120)
    set_repeat_table_header(table.rows[0])

    for number, manual_name, capture in rows:
        row = table.add_row().cells
        values = [number, manual_name, capture[0], capture[1]]
        for idx, cell in enumerate(row):
            cell.width = widths[idx]
            cell.text = values[idx]
            set_cell_border(cell, "DDD8FF")
            set_table_cell_margins(cell, 90, 120, 90, 120)

    path = OUTPUT_DIR / filename
    document.save(path)
    return path


USER_SECTIONS = [
    ManualSection(
        "1. Purpose of This Manual",
        [
            "This manual is for guest users and registered learners who need a complete guide to the learner side of Code Recall.",
            "It explains what each major part of the system does, how to use it, what the current rules mean, and what to do when something feels confusing."
        ],
    ),
    ManualSection(
        "2. What Code Recall Is",
        [
            "Code Recall is a gamified learning platform for Computer System Servicing learners. It combines guided content, structured assessment, review tools, retention flashcards, certificates, and support messaging inside one flow.",
            "The learner side of the system is designed to reward progress, detect weak memory, and guide the learner back to the right part of the subject without forcing a full restart every time."
        ],
        bullets=[
            "Subject-based learning flow",
            "Pre-tests and post-tests",
            "Guided modules",
            "Quiz levels with confidence-based answering",
            "Wrong-answer review",
            "Memory flashcards",
            "XP, levels, and progress tracking",
            "Certificate vault",
            "Contact and support messaging",
        ],
        screenshot_name="user-01-dashboard-overview.png",
        screenshot_description="Full learner dashboard showing XP, progress, subjects, and review areas.",
    ),
    ManualSection(
        "3. User Types",
        [
            "Guest users can explore the system without creating a permanent account. Their progress is temporary and stays only on the current device.",
            "Registered learners can save progress permanently, unlock certificates, use support messaging, and continue review and retention work over time."
        ],
        bullets=[
            "Guest mode is useful for trying the system quickly",
            "Registered accounts are recommended for real progress tracking",
            "Guest progress can be lost after logout, local reset, or browser cleanup",
        ],
        screenshot_name="user-02-login-page.png",
        screenshot_description="Login page showing email login, Google login, and guest mode.",
    ),
    ManualSection(
        "4. How to Start Using the System",
        [
            "A new learner should first sign in or choose guest mode, then open the dashboard and select a subject.",
            "From there, the safest and most effective way to use the system is to follow the subject flow in order instead of jumping directly into later steps."
        ],
        numbered=[
            "Open the login page",
            "Choose guest mode or sign in with an account",
            "Open the dashboard",
            "Select a subject",
            "Complete the subject in order starting with the pre-test",
        ],
    ),
    ManualSection(
        "5. Main Learning Flow",
        [
            "Each subject follows a structured sequence so the learner can move from readiness checking to final completion in a guided way.",
            "The expected order is simple, but each step has a distinct purpose and should not be skipped if the learner wants the best progress and certificate results."
        ],
        numbered=["Pre-Test", "Modules", "Quiz Levels", "Post-Test", "Certificate Unlock"],
        screenshot_name="user-03-subject-flow.png",
        screenshot_description="Subject page with pre-test, modules, quiz levels, post-test, and certificate area visible.",
    ),
    ManualSection(
        "6. Dashboard",
        [
            "The dashboard is the learner's main home page. It shows XP, level, progress, subject shortcuts, continue-learning panels, review tools, due memory-review cards, and support indicators.",
            "Learners should use the dashboard as their daily return point because it highlights what to continue next and what still needs review."
        ],
        screenshot_name="user-04-dashboard-full.png",
        screenshot_description="Full dashboard view with all major learner panels visible.",
    ),
    ManualSection(
        "7. Subjects and Subject Status",
        [
            "Each subject page summarizes the learner's current state in that subject. It shows whether the learner still needs the pre-test, reading modules, quiz levels, post-test, or certificate step.",
            "Subject pages also show assessment bars, so learners can see score and XP summaries for pre-test, quiz track, and post-test."
        ],
    ),
    ManualSection(
        "8. Modules",
        [
            "Modules contain the actual lesson content. They use a guided reading flow with a progress rail, checkpoints, quick checks, and a next-step panel so learners always know where they are in the lesson.",
            "The goal of the module page is not only to show text, but to keep the learner moving in a clear sequence from reading to understanding to practice."
        ],
        screenshot_name="user-05-module-page.png",
        screenshot_description="Module page with lesson content, progress rail, quick check, and next-step areas.",
    ),
    ManualSection(
        "9. Assessments",
        [
            "Pre-tests measure starting knowledge before the learner studies.",
            "Quiz levels reinforce lessons in smaller chunks and are usually completed after modules.",
            "Post-tests confirm what the learner understands after the full learning flow."
        ],
        bullets=[
            "Pre-test: readiness and baseline score",
            "Quiz levels: guided reinforcement and XP opportunities",
            "Post-test: final subject mastery check",
        ],
        screenshot_name="user-06-quiz-page.png",
        screenshot_description="Quiz page showing question, choices, and confidence section.",
    ),
    ManualSection(
        "10. Confidence-Based Answering",
        [
            "After selecting an answer, learners also choose Sure, Somewhat Sure, or Guessing. This allows the system to detect weak memory even when an answer is technically correct.",
            "Confidence choice affects review behavior, flashcard creation, and how the system interprets learner certainty."
        ],
        bullets=[
            "Sure + correct does not create a new flashcard",
            "Somewhat Sure + correct may create a retention card",
            "Guessing + correct may create a retention card",
            "Wrong answers may create review and retention items",
        ],
        screenshot_name="user-07-quiz-level-page.png",
        screenshot_description="Quiz-level page with confidence buttons clearly visible.",
    ),
    ManualSection(
        "11. Retry and Lock Behavior",
        [
            "Some quiz-level items use controlled retry behavior. A learner can retry the same question up to 3 times when that rule is active.",
            "If the learner still fails after the final allowed try, the question may be delayed until the next day. The flashcard and review system then helps the learner return to that concept later."
        ],
    ),
    ManualSection(
        "12. Wrong-Answer Review",
        [
            "Wrong-answer review helps learners revisit missed questions, read the rationale, and reopen the original source activity.",
            "This allows learners to focus directly on weak areas instead of restarting an entire subject or module."
        ],
        screenshot_name="user-08-wrong-answer-review.png",
        screenshot_description="Wrong-answer review page with missed item and rationale.",
    ),
    ManualSection(
        "13. Memory Flashcards",
        [
            "Memory flashcards are used for retention practice. They can be created from wrong answers or low-confidence correct answers, and they can include an image when the original question depends on a picture.",
            "Flashcards are designed to push the learner toward active recall instead of passive rereading."
        ],
        bullets=[
            "Need Again",
            "Hard Recall",
            "Easy Recall",
            "Open Source",
        ],
        screenshot_name="user-09-memory-flashcards.png",
        screenshot_description="Flashcard review page with question, optional image, and recall action buttons.",
    ),
    ManualSection(
        "14. Flashcard Popup Meaning",
        [
            "A flashcard popup can appear for two different reasons. It can appear because the answer just given created a new review item, or because the learner already had due memory cards from earlier work.",
            "This means a popup after a correct answer does not always mean the current answer created a new flashcard. It may simply be reminding the learner about older due cards."
        ],
    ),
    ManualSection(
        "15. Recovery Feedback",
        [
            "When previously missed questions are corrected, the result modal can show a recovery summary so the learner sees that improvement matters.",
            "This helps turn retakes into meaningful recovery instead of making the learner feel that only raw score counts."
        ],
        screenshot_name="user-10-recovery-summary.png",
        screenshot_description="Result modal showing score, XP, and recovery feedback.",
    ),
    ManualSection(
        "16. XP and Anti-Farming Rules",
        [
            "XP is awarded only when a question becomes correct for the first time. Replaying already-correct questions does not generate extra XP.",
            "If a learner was wrong before and later fixes that question, XP is awarded at the moment it becomes correct. This protects the system from XP farming while still rewarding recovery."
        ],
    ),
    ManualSection(
        "17. Certificates and Certificate Vault",
        [
            "Certificates unlock after completing subject requirements. The vault lets learners search, filter, sort, and review issued certificates in a cleaner archive-style page.",
            "Learners can open the certificate preview, download certificates, and review issued-history entries for already unlocked subjects."
        ],
        screenshot_name="user-11-certificate-vault.png",
        screenshot_description="Certificate vault with filters, sorting, and issued-history section.",
    ),
    ManualSection(
        "18. Contact Us",
        [
            "The Contact Us page allows learners to send concerns, comments, feedback, or learning-related questions. Learners can then read admin replies in their own private message history.",
            "Only the learner and authorized admin side should see that private conversation."
        ],
        screenshot_name="user-13-contact-page.png",
        screenshot_description="Learner contact page with message form and private conversation history.",
    ),
    ManualSection(
        "19. Settings",
        [
            "The Settings page includes profile controls, theme preferences, memory review schedule controls, and system actions for safe cleanup or reset.",
            "It is also the place where the learner can export a CSV progress report for reference or teacher review."
        ],
        bullets=[
            "Logout",
            "Reset Hardware Subject",
            "Reset Electrical Subject",
            "Clear Memory Review",
            "Clear Wrong-Answer Review",
            "Reset Progress",
            "Export Progress Report (CSV)",
            "Reset Contact Alerts",
            "Refresh Role Access",
            "Clear Local Data",
        ],
        screenshot_name="user-14-settings-page.png",
        screenshot_description="Settings page with memory review schedule and system actions visible.",
    ),
    ManualSection(
        "20. How to Use the System Day by Day",
        [
            "Code Recall works best when learners use it as a guided cycle instead of skipping randomly from page to page. The dashboard is the best daily starting point because it reflects unfinished subject flow, due memory review, and current support activity in one place.",
            "A learner who follows the same simple routine each session usually gets better retention results than a learner who repeatedly jumps into tests without checking review pressure or incomplete subject steps."
        ],
        numbered=[
            "Start on the dashboard and identify the next unfinished subject task.",
            "Finish one main learning task first, such as a module or quiz level.",
            "Check whether memory review is due before starting another fresh assessment.",
            "Use wrong-answer review when you want deeper correction on specific mistakes.",
            "Return to the dashboard and repeat the cycle instead of opening random pages.",
        ],
    ),
    ManualSection(
        "21. How Scores, Progress Bars, and XP Bars Work",
        [
            "The system now shows several kinds of learner progress, so it is important to understand what each one means. Raw score shows how many questions were answered correctly, while XP shows the reward side of progress under the anti-farming rule.",
            "Assessment bars on the dashboard and subject pages should be read together. A learner can have a decent score but still produce retention work if the answers were low-confidence or if weak topics remain in the queue."
        ],
        bullets=[
            "Score bar: correctness out of the full item total.",
            "XP bar: reward XP earned through first-correct outcomes.",
            "Pre-test: baseline before study.",
            "Quiz Track: cumulative guided practice within the subject.",
            "Post-test: final subject assessment after the learning flow.",
        ],
    ),
    ManualSection(
        "22. What Each Flashcard Button Means in Practice",
        [
            "Flashcard actions are not cosmetic. They tell the retention system how strongly the learner remembered the concept and how the next review should behave.",
            "Learners should answer these honestly. Pressing a stronger recall button than the learner truly earned makes the review schedule less useful later."
        ],
        bullets=[
            "Need Again: the concept still feels weak and should stay under stronger review pressure.",
            "Hard Recall: the learner remembered it, but with effort, so the system advances more carefully.",
            "Easy Recall: the learner remembered it confidently, so the next review can move farther out.",
            "Open Source: reopen the original quiz, module, or question context for deeper understanding.",
        ],
    ),
    ManualSection(
        "23. What the CSV Progress Report Is For",
        [
            "The CSV progress report is a spreadsheet-friendly export of the learner's current system state. It is useful when a learner wants to keep a record, share progress with a teacher, or inspect current values before using reset actions.",
            "It is not a certificate file and not a full visual report of every page. It is a structured data summary designed to open directly in Excel or similar tools."
        ],
        bullets=[
            "Includes summary, progress, results, and schedule-related information.",
            "Can be opened in Excel for quick inspection.",
            "Useful before resets, troubleshooting, or teacher review.",
        ],
    ),
    ManualSection(
        "24. What Each Reset or Clear Action Affects",
        [
            "Not every reset action deletes the same kind of data. Some actions clear subject learning progress, while others only clear review queues or local alerts.",
            "Learners should read the button labels carefully before using them so they do not wipe more than they intended."
        ],
        bullets=[
            "Reset Hardware Subject: clears Computer Hardware learning progress.",
            "Reset Electrical Subject: clears Electrical learning progress.",
            "Clear Memory Review: removes retention flashcards without fully wiping subject completion.",
            "Clear Wrong-Answer Review: removes the wrong-answer backlog only.",
            "Reset Progress: broader progress reset across learner data.",
            "Reset Contact Alerts: clears local reply-alert markers only.",
            "Clear Local Data: clears device-side state and should be used carefully, especially for guest sessions.",
        ],
    ),
    ManualSection(
        "25. Guest and Registered Data Behavior",
        [
            "Guest mode is intentionally temporary and depends on the current browser and device. Registered accounts provide stronger continuity because their progress and role-sensitive features are tied to the account rather than only the local browser state.",
            "This difference matters when using resets, switching devices, or trying to keep a long-term learning record."
        ],
        bullets=[
            "Guest progress can disappear after logout, browser cleanup, or local resets.",
            "Registered learners should use sign-in mode for long-term study tracking.",
            "Private contact history and privileged navigation depend on the current signed-in account state.",
        ],
    ),
    ManualSection(
        "26. Troubleshooting Highlights",
        [
            "If flashcards appear often, review confidence choices, wrong answers, and due cards from earlier work.",
            "If XP does not increase after replaying, the learner may have already earned first-correct XP for those questions.",
            "If a certificate is still locked, confirm that pre-test, modules, quiz levels, and post-test were completed.",
            "If guest progress disappears, remember that guest mode is temporary and device-based."
        ],
    ),
]


ADMIN_SECTIONS = [
    ManualSection(
        "1. Purpose of This Manual",
        [
            "This manual is for admins such as teachers, facilitators, and school staff assigned to monitor learners.",
            "It explains how to use the admin side of Code Recall to review learner progress, retention, support messages, and protected admin access."
        ],
    ),
    ManualSection(
        "2. Admin Role Overview",
        [
            "Admins monitor learner performance, contact messages, weak topics, and retention load. Their job is to identify learners who need intervention and help keep support communication organized.",
            "Admins should read the dashboard not only as a score report, but also as an indicator of learner memory stability and review behavior."
        ],
    ),
    ManualSection(
        "3. Admin Access and 2FA",
        [
            "Admins now use authenticator-based 2FA before entering the admin dashboard. Setup includes QR enrollment, a manual key, direct authenticator launch, backup codes, and final verification.",
            "Admins should save backup codes during enrollment and use them only for recovery."
        ],
        screenshot_name="admin-01-admin-mfa-setup.png",
        screenshot_description="Admin MFA setup page showing QR, manual key, and backup codes.",
    ),
    ManualSection(
        "4. Entering the Admin Dashboard",
        [
            "After successful login and 2FA verification, the admin is routed into the admin dashboard.",
            "This dashboard is the main working view for teachers or facilitators. It should be checked regularly to spot learners who are behind, topics that are repeatedly missed, and retention load that is getting too high."
        ],
    ),
    ManualSection(
        "5. Admin Dashboard",
        [
            "The admin dashboard provides summary visibility into learner performance and support needs."
        ],
        bullets=[
            "Average pre-test score",
            "Average post-test score",
            "Hardware completion",
            "Electrical completion",
            "Most-missed topics",
            "Retention indicators",
        ],
        screenshot_name="admin-02-dashboard-overview.png",
        screenshot_description="Admin dashboard with overview cards and learner monitoring panels.",
    ),
    ManualSection(
        "6. Learner Analytics",
        [
            "The learner profile modal provides detailed analytics for one learner. This is the best place to understand whether the learner is struggling with content, confidence, or retention.",
            "Admins should use this view before replying to learning concerns because it provides stronger context than the overview cards alone."
        ],
        bullets=[
            "Assessment Performance",
            "Weak Topics",
            "Recent Activity History",
            "Retention Snapshot",
            "Due Memory Queue",
        ],
        screenshot_name="admin-03-learner-modal.png",
        screenshot_description="Learner modal showing assessment bars, weak topics, and retention panels.",
    ),
    ManualSection(
        "7. How to Read the Assessment Performance Panels",
        [
            "Assessment bars help compare readiness, practice performance, and final mastery. A learner with a strong post-test but many weak topics may still need review support, while a learner with a weak pre-test and stronger later results may simply be progressing normally.",
            "Admins should compare score bars with retention indicators instead of reading only one number in isolation."
        ],
        bullets=[
            "Low pre-test plus stronger post-test often means the learning path is helping.",
            "Weak quiz track with completed modules may suggest passive reading without strong retrieval.",
            "Strong score plus heavy due-card load can still mean fragile memory.",
            "Retention recovery matters because it shows whether previous weak areas are being repaired.",
        ],
    ),
    ManualSection(
        "8. Retention Analytics",
        [
            "Admins can monitor retention load across the learner base using the retention indicators on the dashboard.",
            "These indicators show whether learners are actually returning to weak concepts and whether the memory-review system is becoming too heavy or being ignored."
        ],
        bullets=[
            "Due Memory Cards",
            "Learners With Due Cards",
            "Low-confidence cards",
            "Retention recoveries",
            "Highest due flashcard load",
            "Lowest retention recovery",
        ],
        screenshot_name="admin-05-retention-indicators.png",
        screenshot_description="Admin dashboard area showing due-card and retention indicators.",
    ),
    ManualSection(
        "9. How to Use the Learner Profile Modal",
        [
            "The learner profile modal is the best intervention screen on the admin side because it combines performance, weak topics, recent activity, and retention state in one place.",
            "Admins should open this modal whenever a learner reports confusion, when the dashboard shows unusual retention pressure, or when a learner's scores do not match their visible progress."
        ],
        numbered=[
            "Open the learner from the admin list or analytics area.",
            "Check assessment performance first to see whether the issue is broad or topic-specific.",
            "Review weak topics and recent activity for patterns, not only isolated mistakes.",
            "Check retention snapshot and due memory queue to see whether review pressure is too high.",
            "Use this context before replying to the learner or giving advice.",
        ],
    ),
    ManualSection(
        "10. Weak Topics and Intervention",
        [
            "Most-missed topics and learner weak-topic lists help admins identify where intervention should happen first.",
            "If several learners miss the same concept, the issue may be instructional and not only individual."
        ],
    ),
    ManualSection(
        "11. Contact Inbox",
        [
            "Admins can open learner support threads, read concerns, and reply directly through the admin contact area.",
            "Support messages may include system issues, subject confusion, or learning concerns that need teacher follow-up."
        ],
        screenshot_name="admin-04-contact-inbox.png",
        screenshot_description="Admin contact inbox with learner messages and reply controls.",
    ),
    ManualSection(
        "12. How to Respond to Common Learner Situations",
        [
            "Admin work becomes much easier when the numbers are translated into clear guidance for the learner. Most intervention does not require a technical fix; it requires pointing the learner to the right next action.",
            "The system provides enough signals to recommend whether the learner should continue modules, revisit wrong answers, reduce pace, or complete memory review first."
        ],
        bullets=[
            "High due-card load: ask the learner to complete memory review before opening more new quizzes.",
            "Many low-confidence correct answers: explain that the learner should slow down and use flashcards honestly.",
            "Weak topics concentrated in one subject: direct the learner back to that subject's review or source material.",
            "No XP gain complaint: explain first-correct-only XP and recovery-based earning.",
            "Low completion with little recent activity: encourage returning to the dashboard and continuing the subject flow in order.",
        ],
    ),
    ManualSection(
        "13. Privacy and Conduct",
        [
            "Learner conversation history is intended to stay private between the learner and the authorized admin side.",
            "Admins should avoid exposing one learner's thread to another learner and should use the inbox only for legitimate support and monitoring."
        ],
    ),
    ManualSection(
        "14. Suggested Admin Routine",
        ["A suggested daily or weekly routine keeps the admin side manageable and useful."],
        numbered=[
            "Review Contact Us inbox",
            "Check learners with high due-card load",
            "Check most-missed topics",
            "Review completion rates",
            "Open learner analytics for struggling users",
            "Reply to urgent learner concerns",
        ],
    ),
    ManualSection(
        "15. Admin 2FA Recovery",
        [
            "Admins who lose access to their authenticator app should first try an unused backup code.",
            "If backup codes are unavailable or exhausted, the admin can use Reset My Admin 2FA from the admin page and complete re-enrollment."
        ],
    ),
    ManualSection(
        "16. Troubleshooting Highlights",
        [
            "If a learner says they earned no XP, confirm whether first-correct-only XP prevented farming.",
            "If an admin cannot enter the dashboard, confirm admin 2FA enrollment, the current authenticator code, or backup-code availability.",
            "If dashboard numbers look unusual, confirm that results and retention data were saved correctly.",
            "If contact threads look incomplete, confirm the learner actually sent the message from the current account."
        ],
    ),
]


SUPERADMIN_SECTIONS = [
    ManualSection(
        "1. Purpose of This Manual",
        [
            "This manual is for super admins and system managers who operate the platform at the highest privilege level.",
            "It focuses on oversight, security, privileged access control, and safe system operation rather than on code-level implementation."
        ],
    ),
    ManualSection(
        "2. Super Admin Role Overview",
        [
            "Super admins handle system oversight, security monitoring, privileged-role management, retention oversight, and operational validation.",
            "They do not only monitor learner performance. They also monitor whether the platform's controls, roles, security, and oversight flows are working correctly."
        ],
    ),
    ManualSection(
        "3. Dashboard Purpose",
        [
            "The super-admin dashboard provides system-wide oversight beyond the normal admin view. It combines learner analytics with privileged-role monitoring and access control.",
            "This page should be treated as a control room for the platform, not just another teacher dashboard."
        ],
        screenshot_name="superadmin-01-dashboard-overview.png",
        screenshot_description="Super-admin dashboard with oversight, analytics, and privileged controls.",
    ),
    ManualSection(
        "4. Super Admin 2FA",
        [
            "Super admins now use authenticator-based 2FA before they can open the super-admin dashboard. The setup flow includes a locally generated QR code, manual setup key, direct authenticator launch, backup codes, and verification before access.",
            "Because super-admin accounts have the highest privilege, this security step should never be skipped or shared."
        ],
        bullets=[
            "QR code enrollment",
            "Manual setup key",
            "Open in Authenticator",
            "Backup codes",
            "Verification before access",
            "Reset My Super Admin 2FA",
        ],
        screenshot_name="superadmin-02-super-admin-mfa-setup.png",
        screenshot_description="Super-admin MFA setup page with local QR generation and backup-code tools.",
    ),
    ManualSection(
        "5. 2FA Oversight",
        [
            "The 2FA Oversight section tracks the protection status of privileged accounts and helps confirm whether admin and super-admin accounts are actually protected.",
            "This turns 2FA into an actively monitored system control instead of a hidden sign-in requirement."
        ],
        bullets=[
            "Privileged accounts",
            "2FA enrolled",
            "Pending enrollment",
            "Verified today",
            "Low backup reserves",
            "Backup-code sign-ins",
            "Recovery Risk Accounts",
        ],
        screenshot_name="superadmin-03-2fa-oversight.png",
        screenshot_description="2FA Oversight panel showing enrollment, verification, and recovery-risk indicators.",
    ),
    ManualSection(
        "6. How to Read 2FA Oversight",
        [
            "A privileged account can be enrolled and still be risky if backup reserves are almost gone or if the user relies too often on backup-code sign-ins. The oversight panel should be treated as a health monitor, not just a simple enabled-or-disabled list.",
            "Super admins should use this area to detect future lockout risk before a privileged account actually loses access."
        ],
        bullets=[
            "Pending enrollment: privileged user still has access but has not completed MFA setup.",
            "Verified today: confirms recent privileged verification activity succeeded.",
            "Low backup reserves: recovery safety is weakening and re-enrollment may be needed soon.",
            "Backup-code sign-ins: should be occasional and reviewed if they become frequent.",
            "Recovery risk accounts: accounts that may need intervention before access becomes unstable.",
        ],
    ),
    ManualSection(
        "7. 2FA Recovery Policy",
        [
            "Backup codes should be treated as emergency recovery only. Super admins should store them offline, prefer normal authenticator sign-ins, and re-enroll before backup reserves become too low.",
            "Any account repeatedly using backup-code sign-ins should be reviewed because that may indicate a risky recovery habit or poor device control."
        ],
    ),
    ManualSection(
        "8. Email Access Control",
        [
            "Email Access Control allows super admins to grant or remove admin and super-admin access by email without editing project files directly.",
            "This is powerful and should be used carefully because it changes real access to protected parts of the system."
        ],
        bullets=[
            "Grant admin access",
            "Grant super_admin access",
            "Remove granted access with strong warning",
            "Extra-strong warning for removing current active super-admin access",
        ],
        screenshot_name="superadmin-04-email-access-control.png",
        screenshot_description="Granted Email Access and access-control form in the super-admin page.",
    ),
    ManualSection(
        "9. How to Safely Change Privileged Access",
        [
            "Access changes should be treated as security-sensitive actions, not casual settings. Removing or changing the wrong privileged grant can lock out an operator or leave the system without a safe recovery path.",
            "The system already shows stronger warnings, especially when the current active super-admin grant is being removed, but super admins should still confirm the broader operational situation first."
        ],
        numbered=[
            "Confirm which account is currently signed in.",
            "Confirm whether another super admin remains available as fallback.",
            "Review whether the target account has already enrolled in 2FA.",
            "Proceed with the grant or removal only after those checks.",
            "Recheck access behavior after the change by refreshing role access or re-logging in.",
        ],
    ),
    ManualSection(
        "10. Learner and Retention Oversight",
        [
            "Super admins should monitor whether the retention queue is too aggressive, whether due-card load is too high, and whether recovery is actually happening across the learner base.",
            "If many learners keep building due-card pressure without recovery, the schedule or learning flow may need adjustment."
        ],
    ),
    ManualSection(
        "11. Recovery Management",
        [
            "The super-admin role includes responsibility for safe recovery policy. A strong recovery setup means there is still a path back into the platform if one privileged user loses a device, uses too many backup codes, or resets their 2FA incorrectly.",
            "Super admins should periodically review backup reserve warnings and decide when re-enrollment is safer than continuing to operate on low recovery margin."
        ],
        bullets=[
            "Backup codes are emergency-only, not a normal sign-in method.",
            "Repeated backup-code use should prompt investigation or re-enrollment.",
            "Reset and re-enroll should be done carefully, especially if few privileged accounts remain.",
        ],
    ),
    ManualSection(
        "12. System Configuration Review",
        [
            "Super admins should be familiar with the system's retention schedules, first-correct-only XP, certificate logic, contact privacy rules, role-based navigation, and access syncing behavior.",
            "This is important because many issues reported by admins or learners come from interactions between these systems rather than from one isolated feature."
        ],
    ),
    ManualSection(
        "13. Operational Safety Practices",
        [
            "Before changing privileged access, confirm which account is currently logged in and whether that account is still the primary recovery path.",
            "Before resetting 2FA or removing access, confirm whether backup codes were saved and whether another privileged account can still manage the system if needed."
        ],
    ),
    ManualSection(
        "14. Monitoring Routine",
        ["A recommended super-admin routine helps keep the platform healthy."],
        numbered=[
            "Check 2FA Oversight for pending or risky privileged accounts",
            "Check learner retention pressure and recovery indicators",
            "Review Email Access Control for outdated or risky grants",
            "Spot-check contact inbox activity and admin response flow",
            "Confirm recent code or rule changes did not break protected workflows",
        ],
    ),
    ManualSection(
        "15. Troubleshooting Highlights",
        [
            "If 2FA oversight looks incomplete, confirm the latest Firestore rules are deployed and that privileged users have enrolled.",
            "If email access actions fail, confirm the current account still has super-admin rights.",
            "If retention analytics look wrong, inspect learner result data, wrong-answer review, and retention queue payloads.",
            "If a super admin is stuck on verification, confirm that the current QR or manual key was used and that backup codes are still valid."
        ],
    ),
]


DEVELOPER_SECTIONS = [
    ManualSection(
        "1. Purpose of the Developer Manual",
        [
            "This manual is for developers and maintainers who need to understand how the current Code Recall project is structured, what the main systems do, and where to look when changing or debugging the platform.",
            "It is separate from the super-admin manual so technical implementation details can live in one place without overloading the operational guide."
        ],
    ),
    ManualSection(
        "2. Current Stack",
        [
            "Code Recall is currently a multi-page web application using plain HTML, CSS, and vanilla JavaScript modules. Authentication and persisted data are handled through Firebase Authentication and Firestore.",
            "The current architecture is page-based rather than framework-based. Each route has its own HTML page, CSS file, and script file or module cluster."
        ],
        bullets=[
            "HTML page per route",
            "Dedicated CSS per major page or feature area",
            "Vanilla JavaScript modules and page scripts",
            "Firebase Authentication for sign-in and role-aware routing",
            "Firestore for remote data and rules enforcement",
            "Local storage and session storage for device-side state",
        ],
    ),
    ManualSection(
        "3. Major Functional Areas",
        [
            "The project currently includes several major systems that interact closely. Developers should understand these relationships before refactoring or extending features."
        ],
        bullets=[
            "Learner dashboard and subject progression",
            "Module reading flow and checkpoint guidance",
            "Quiz and quiz-level assessment flows",
            "Wrong-answer review and retention queue",
            "Memory flashcards and spaced review",
            "Certificates and certificate vault",
            "Contact and support messaging",
            "Admin analytics and learner monitoring",
            "Super-admin access control and 2FA oversight",
        ],
        screenshot_name="developer-01-system-scope.png",
        screenshot_description="Optional system map or architecture illustration if you create one.",
    ),
    ManualSection(
        "4. How the Application Is Organized",
        [
            "Code Recall is organized as a multi-page application rather than a single framework-driven frontend. Most routes have their own HTML page, CSS file, and one or more supporting scripts, with shared helpers used for common state or auth behavior.",
            "This structure keeps individual pages understandable, but it also means repeated logic can drift if developers update one flow and forget the others."
        ],
        bullets=[
            "Public pages for marketing, landing, and general information.",
            "Learner pages for dashboard, subjects, modules, quizzes, review, certificates, contact, and settings.",
            "Privileged pages for admin, super admin, and MFA verification or enrollment.",
            "Shared helper scripts for auth, retention, QR generation, session verification, and exports.",
        ],
    ),
    ManualSection(
        "5. Data and State Layers",
        [
            "The system uses several storage layers, and many bugs come from misunderstanding which layer is responsible for a specific behavior. Developers should identify whether the issue starts in Firestore, localStorage, sessionStorage, or derived UI state before patching symptoms.",
            "In practice, state often flows from storage into computed summaries and then into role-aware UI. That means a single stale value can surface as a dashboard issue, a review issue, or a privileged-access issue depending on where it is consumed."
        ],
        bullets=[
            "Firestore: persistent account, progress, result, security, contact, and analytics records.",
            "localStorage: guest state, device-side progress helpers, retention schedule preferences, and export-related cached values.",
            "sessionStorage: current privileged verification state and short-lived session gates.",
            "Derived UI state: dashboard summaries, bars, warnings, and next-step guidance built from stored data.",
        ],
    ),
    ManualSection(
        "6. Learning and Progress Model",
        [
            "The learner flow is built around pre-test, modules, quiz levels, post-test, and certificate unlock. Progress is surfaced on the dashboard, subject pages, and certificate vault.",
            "When changing any one of these areas, developers should check whether the same progress state is also used by analytics, certificate eligibility, retention creation, or review flows."
        ],
    ),
    ManualSection(
        "7. XP and Anti-Farming Model",
        [
            "The current XP rule is first-correct-only. Re-answering a question that was already correct should not produce more XP. If a learner was previously wrong and later fixes that item, XP is awarded at the time it becomes correct.",
            "This rule exists in both the quiz and quiz-level flows, so changes to scoring must be reviewed in both areas to avoid reintroducing XP farming."
        ],
    ),
    ManualSection(
        "8. Retention and Flashcard System",
        [
            "The retention system uses wrong answers and low-confidence correct answers to seed memory-review items. Those items can become due immediately or later based on the configured retention schedule.",
            "The dashboard, review page, settings page, and admin analytics all depend on this retention data, so retention changes usually affect more than one page."
        ],
        bullets=[
            "Wrong answers can enter review and retention",
            "Low-confidence correct answers can enter retention",
            "Need Again, Hard Recall, and Easy Recall affect scheduling behavior",
            "Settings can change the memory review schedule",
            "Admin and super-admin analytics read retention pressure and recovery data",
        ],
    ),
    ManualSection(
        "9. How the Retention Pipeline Works",
        [
            "The retention pipeline starts at answer evaluation, not at the flashcard screen. Once a learner answers, the system can seed wrong-answer review, seed or update retention items, and recalculate due review state for both the learner and the admin analytics view.",
            "When debugging memory-review behavior, developers should trace the entire lifecycle from answer event to stored state to dashboard counts to review-page rendering."
        ],
        numbered=[
            "Learner answers a question.",
            "The system checks correctness and confidence.",
            "Wrong-answer review and retention queue are updated if conditions are met.",
            "Due review counts and flashcard availability are recalculated.",
            "Admin and super-admin analytics consume the resulting state.",
        ],
    ),
    ManualSection(
        "10. Contact System",
        [
            "The contact system is role-sensitive. Learners should only see their own thread history. Admins and super admins can read and reply through the authorized inbox views.",
            "When modifying contact logic, developers should test both learner privacy and reply visibility carefully."
        ],
    ),
    ManualSection(
        "11. 2FA Architecture",
        [
            "Privileged 2FA is now implemented for super admins and admins. Each role has its own setup and verification page, while both store security profile data in Firestore.",
            "The setup flow includes QR or manual enrollment, backup codes, verification, session marking, and dashboard or page guards."
        ],
        bullets=[
            "super-admin-mfa.html and admin-mfa.html handle enrollment and verification",
            "securityProfiles stores MFA enrollment and verification state",
            "backupCodeUseCount and lastVerificationMethod support oversight",
            "local QR generation avoids external QR-service dependency",
            "dashboard and protected-page guards enforce role-specific 2FA verification",
        ],
        screenshot_name="developer-02-mfa-flow.png",
        screenshot_description="Optional capture of MFA setup page or technical flow diagram.",
    ),
    ManualSection(
        "12. Page Guards, Role Routing, and Access Refresh",
        [
            "Role-sensitive behavior now depends on more than one check. A user may be authenticated but still redirected because of role mismatch, missing privileged 2FA verification, or stale role state.",
            "This is why some navigation issues are actually auth or storage issues. Developers should inspect both route logic and current role session state before assuming a pure UI bug."
        ],
    ),
    ManualSection(
        "13. Firestore Rules Awareness",
        [
            "Many bugs in this project are not UI bugs but permission bugs. Role-based features such as contact, security profiles, access control, and analytics depend on Firestore rules being aligned with the scripts that call them.",
            "Whenever a feature adds a new collection, write shape, or oversight read path, developers should review whether rules must also be updated and redeployed."
        ],
    ),
    ManualSection(
        "14. Settings and Local State",
        [
            "The Settings page includes actions that affect different layers of state: subject resets, memory review clearing, wrong-answer clearing, CSV export, contact alert reset, role refresh, and local data clearing.",
            "Because the app uses both local and remote state, developers should be careful to clear or preserve only the intended data when adding new reset actions."
        ],
    ),
    ManualSection(
        "15. Export and Reporting Considerations",
        [
            "The current CSV export is intentionally simple and spreadsheet-friendly. It works well as a system snapshot, but it is not yet a polished multi-sheet reporting workflow.",
            "If export requirements grow later, developers should decide early whether the target audience is learners, teachers, or technical maintainers, because those audiences need different report structures."
        ],
    ),
    ManualSection(
        "16. Recommended Test Pass After Feature Changes",
        ["A practical post-change verification pass should cover the main dependent systems, not just the area directly edited."],
        numbered=[
            "Open learner dashboard and subject pages",
            "Run at least one quiz and one quiz-level flow",
            "Check confidence-based flashcard behavior",
            "Check wrong-answer review and memory flashcards",
            "Verify XP does not farm on already-correct retakes",
            "Verify certificate unlock or vault display if progress logic changed",
            "Check admin analytics if learner results or retention changed",
            "Check super-admin 2FA and oversight if security flows changed",
            "Check contact privacy if message or role logic changed",
        ],
    ),
    ManualSection(
        "17. Debugging Habits That Fit This Project",
        [
            "Because Code Recall mixes page-local logic with shared storage and role-aware behavior, the best debugging pattern is to identify which layer is failing first: rendering, local state, session state, remote data, or rules enforcement.",
            "Developers should avoid patching only the final visible symptom. Many issues that appear on one page are caused by stale cached values, missing Firestore rule deploys, or helper behavior shared by several pages."
        ],
        bullets=[
            "Check console errors first.",
            "Check local and session storage next.",
            "Check Firestore rule assumptions after that.",
            "Then inspect page-specific rendering or event handling.",
        ],
    ),
    ManualSection(
        "18. Suggested Refactor Priorities",
        [
            "If the project keeps growing, the most helpful refactor path is not a random framework migration. It is better to first improve structure around shared state, repeated UI logic, and cross-page helpers.",
            "The most repeated or sensitive systems right now are retention, scoring, role checks, contact state, and privileged security flows."
        ],
    ),
    ManualSection(
        "19. Future Technical Documentation",
        [
            "A future deeper technical manual can include collection maps, storage keys, release workflows, page ownership, and deployment checklists. This file is meant to be the readable starting point before that deeper reference exists."
        ],
    ),
]


CHECKLIST_ROWS = [
    ("1", "User Manual", ("user-01-dashboard-overview.png", "Full learner dashboard overview")),
    ("2", "User Manual", ("user-02-login-page.png", "Login page with guest and registered options")),
    ("3", "User Manual", ("user-03-subject-flow.png", "Subject flow page with pre-test, modules, quiz levels, post-test, certificate")),
    ("4", "User Manual", ("user-04-dashboard-full.png", "Full dashboard view")),
    ("5", "User Manual", ("user-05-module-page.png", "Module page with progress rail and content")),
    ("6", "User Manual", ("user-06-quiz-page.png", "Quiz page with answers and confidence")),
    ("7", "User Manual", ("user-07-quiz-level-page.png", "Quiz-level page with confidence section")),
    ("8", "User Manual", ("user-08-wrong-answer-review.png", "Wrong-answer review page")),
    ("9", "User Manual", ("user-09-memory-flashcards.png", "Memory flashcards with recall buttons")),
    ("10", "User Manual", ("user-10-recovery-summary.png", "Result modal with recovery summary")),
    ("11", "User Manual", ("user-11-certificate-vault.png", "Certificate vault page")),
    ("12", "User Manual", ("user-12-certificate-preview.png", "Opened certificate preview")),
    ("13", "User Manual", ("user-13-contact-page.png", "Learner contact page")),
    ("14", "User Manual", ("user-14-settings-page.png", "Settings page with system actions")),
    ("15", "Admin Manual", ("admin-01-admin-mfa-setup.png", "Admin MFA setup page")),
    ("16", "Admin Manual", ("admin-02-dashboard-overview.png", "Admin dashboard overview")),
    ("17", "Admin Manual", ("admin-03-learner-modal.png", "Admin learner analytics modal")),
    ("18", "Admin Manual", ("admin-04-contact-inbox.png", "Admin contact inbox")),
    ("19", "Admin Manual", ("admin-05-retention-indicators.png", "Admin retention indicators section")),
    ("20", "Super Admin Manual", ("superadmin-01-dashboard-overview.png", "Super-admin dashboard overview")),
    ("21", "Super Admin Manual", ("superadmin-02-super-admin-mfa-setup.png", "Super-admin MFA setup page")),
    ("22", "Super Admin Manual", ("superadmin-03-2fa-oversight.png", "2FA Oversight section")),
    ("23", "Super Admin Manual", ("superadmin-04-email-access-control.png", "Email Access Control section")),
    ("24", "Developer Manual", ("developer-01-system-scope.png", "Optional system map or scope image")),
    ("25", "Developer Manual", ("developer-02-mfa-flow.png", "Optional technical or MFA flow image")),
]


def main() -> None:
    build_manual(
        "Code_Recall_User_Manual.docx",
        "Code Recall User Manual",
        "Comprehensive guide for guest users and registered learners",
        "Learners and guest users",
        "Code Recall | User Manual",
        USER_SECTIONS,
    )
    build_manual(
        "Code_Recall_Admin_Manual.docx",
        "Code Recall Admin Manual",
        "Comprehensive guide for teachers, facilitators, and monitoring staff",
        "Admins and teachers",
        "Code Recall | Admin Manual",
        ADMIN_SECTIONS,
    )
    build_manual(
        "Code_Recall_Super_Admin_Manual.docx",
        "Code Recall Super Admin Manual",
        "Operational guide for privileged platform operators",
        "Super admins and system managers",
        "Code Recall | Super Admin Manual",
        SUPERADMIN_SECTIONS,
    )
    build_manual(
        "Code_Recall_Developer_Manual.docx",
        "Code Recall Developer Manual",
        "Technical guide for maintainers and future contributors",
        "Developers and maintainers",
        "Code Recall | Developer Manual",
        DEVELOPER_SECTIONS,
    )
    build_checklist_doc(
        "Code_Recall_Manual_Screenshot_Checklist.docx",
        "Code Recall Manual Screenshot Checklist",
        CHECKLIST_ROWS,
    )


if __name__ == "__main__":
    main()
